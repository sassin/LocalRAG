# rag/index.py
"""
Simple local indexer for "regular" documents under resources/data.

Supported:
- .pdf   (text-based PDFs via pypdf; no OCR; tables may be imperfect)
- .txt, .md
- .docx
- .csv   (basic CSV -> tab-separated text; no pandas)
- .json  (stored as pretty text)

Outputs:
- resources/.rag_store/index.faiss
- resources/.rag_store/records.json
- resources/.rag_store/meta.json

Install:
pip install faiss-cpu sentence-transformers numpy pypdf python-docx
"""

from __future__ import annotations

import csv
import json
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

import numpy as np

try:
    import faiss  # type: ignore
except Exception as e:
    raise RuntimeError("faiss is required. Install with: pip install faiss-cpu") from e

try:
    from sentence_transformers import SentenceTransformer  # type: ignore
except Exception as e:
    raise RuntimeError("sentence-transformers required. Install with: pip install sentence-transformers") from e

try:
    from pypdf import PdfReader  # type: ignore
except Exception as e:
    raise RuntimeError("pypdf required for PDFs. Install with: pip install pypdf") from e

try:
    from docx import Document  # type: ignore
except Exception as e:
    raise RuntimeError("python-docx required. Install with: pip install python-docx") from e


SUPPORTED_EXTS = {".pdf", ".txt", ".md", ".docx", ".csv", ".json"}


@dataclass
class Block:
    page: Optional[int]  # for PDFs only (1-based). else None
    text: str


def _clean(s: str) -> str:
    return (s or "").replace("\r\n", "\n").strip()


def chunk_text(text: str, chunk_size: int = 1400, overlap: int = 250) -> List[str]:
    text = _clean(text)
    if not text:
        return []
    chunks: List[str] = []
    n = len(text)
    start = 0
    while start < n:
        end = min(n, start + chunk_size)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(0, end - overlap)
    return chunks


def discover_docs(root: Path) -> List[Path]:
    root = root.resolve()
    return sorted([p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS])


# -----------------------
# Extractors (minimal)
# -----------------------

def extract_pdf(path: Path) -> List[Block]:
    blocks: List[Block] = []
    reader = PdfReader(str(path))
    for i, page in enumerate(reader.pages):
        try:
            t = _clean(page.extract_text() or "")
        except Exception:
            t = ""
        if t:
            blocks.append(Block(page=i + 1, text=t))
    return blocks


def extract_docx(path: Path) -> List[Block]:
    doc = Document(str(path))
    lines: List[str] = []
    for p in doc.paragraphs:
        t = _clean(p.text)
        if t:
            lines.append(t)

    # Minimal table extraction: join cells with tabs
    for table in doc.tables:
        for row in table.rows:
            cells = [_clean(c.text).replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))

    return [Block(page=None, text="\n".join(lines))]


def extract_textfile(path: Path) -> List[Block]:
    return [Block(page=None, text=path.read_text(encoding="utf-8", errors="ignore"))]


def extract_csv(path: Path, max_rows: int = 3000) -> List[Block]:
    rows: List[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i >= max_rows:
                break
            rows.append("\t".join([_clean(c) for c in row]))
    return [Block(page=None, text="\n".join(rows))]


def extract_json(path: Path) -> List[Block]:
    try:
        obj = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        return [Block(page=None, text=json.dumps(obj, ensure_ascii=False, indent=2))]
    except Exception:
        # fallback to raw
        return [Block(page=None, text=path.read_text(encoding="utf-8", errors="ignore"))]


def extract_document(path: Path) -> List[Block]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in {".txt", ".md"}:
        return extract_textfile(path)
    if ext == ".csv":
        return extract_csv(path)
    if ext == ".json":
        return extract_json(path)
    return [Block(page=None, text="")]


# -----------------------
# Index builder
# -----------------------

def build_index(
    docs_root: Path,
    store_dir: Path,
    model_name: str = "sentence-transformers/all-MiniLM-L6-v2",
) -> None:
    docs_root = docs_root.resolve()
    store_dir = store_dir.resolve()
    store_dir.mkdir(parents=True, exist_ok=True)

    docs = discover_docs(docs_root)
    if not docs:
        raise RuntimeError(f"No supported files found under: {docs_root}")

    embedder = SentenceTransformer(model_name)

    records: List[Dict[str, Any]] = []
    texts: List[str] = []

    num_no_text = 0

    for doc_path in docs:
        rel = doc_path.relative_to(docs_root).as_posix()

        try:
            blocks = extract_document(doc_path)
        except Exception:
            blocks = []

        has_text = False
        for block in blocks:
            for ci, chunk in enumerate(chunk_text(block.text)):
                has_text = True
                records.append(
                    {
                        "source_path": rel,
                        "page": block.page,
                        "chunk_index": ci,
                        "text": chunk,
                    }
                )
                texts.append(chunk)

        if not has_text:
            num_no_text += 1

    if not texts:
        raise RuntimeError("No extractable text found. PDFs may be scanned/image-only (needs OCR in a separate pipeline).")

    X = np.asarray(embedder.encode(texts, normalize_embeddings=True), dtype="float32")
    index = faiss.IndexFlatIP(X.shape[1])
    index.add(X)

    faiss.write_index(index, str(store_dir / "index.faiss"))
    (store_dir / "records.json").write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    (store_dir / "meta.json").write_text(
        json.dumps(
            {
                "model_name": model_name,
                "docs_root": str(docs_root),
                "store_dir": str(store_dir),
                "num_files_found": len(docs),
                "num_files_with_no_text": num_no_text,
                "num_chunks": len(records),
                "supported_exts": sorted(SUPPORTED_EXTS),
                "pdf_extractor": "pypdf",
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    print(f"✅ Indexed {len(records)} chunks from {len(docs)} files")
    print(f"   Store: {store_dir}")


def main():
    base_dir = Path(__file__).resolve().parents[1]
    docs_root = base_dir / "resources" / "data"
    store_dir = base_dir / "resources" / ".rag_store"
    model_name = "sentence-transformers/all-MiniLM-L6-v2"

    args = sys.argv[1:]
    if "--docs" in args:
        docs_root = Path(args[args.index("--docs") + 1])
    if "--store" in args:
        store_dir = Path(args[args.index("--store") + 1])
    if "--model" in args:
        model_name = args[args.index("--model") + 1]

    build_index(docs_root=docs_root, store_dir=store_dir, model_name=model_name)


if __name__ == "__main__":
    main()
